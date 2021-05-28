import requests
from docx import Document
from docx.shared import Inches
import requests.exceptions
from tabulate import tabulate
from netaddr import IPAddress
from netaddr.core import AddrFormatError
def ip_address(ip_name):


    response=requests.get(url="https://www.virustotal.com/vtapi/v2/ip-address/report",
                    params= {"apikey":"6da139d5fcc89d75b423803039344cf46f765d8985ec5c36a4cee49a0068d758","ip":ip_name})
    if response.status_code==200:
        result = response.json()
        try:
            document = Document()
            document.add_heading("virustotal", level=2)
            table = document.add_table(rows=2, cols=2)
            table.style = 'Table Grid'
            row = table.rows[0].cells
            row[0].text = 'ip'
            row[1].text = 'country'
            row = table.rows[1].cells

            row[0].text=[ip_name]
            row[1].text=[result.get("country")]

            #print('\n\n', tabulate(lt2, headers=['ip','country'], tablefmt='orgtbl'))
            document.add_heading("resolutions",level=2)
            tab1=document.add_table(rows=1,cols=2)
            tab1.style='Table Grid'
            row=tab1.rows[0].cells
            row[0].text='last_resolved'
            row[0].width = Inches(2.0)
            row[1].text='hostname'
            try:
                if len(result['resolutions']) != 0:
                    for ele in result['resolutions']:
                        row=tab1.add_row().cells
                        row[0].text=[ele["last_resolved"]]
                        row[0].width = Inches(2.0)
                        row[1].text=([ele["hostname"]])
            except KeyError:
                row = tab1.add_row().cells
                row[0].text=(["no info"])
                row[1].text=(["no info"])
                #print('\n\n', tabulate(lt2, headers=['last_resolved', 'hostname'], tablefmt='orgtbl'))
            document.add_heading("detected_urls",level=2)
            tab2 = document.add_table(rows=1, cols=2)
            tab2.style = 'Table Grid'
            row = tab2.rows[0].cells
            row[0].text = 'url'
            row[1].text = 'score'
            try:
                if len(result['detected_urls'])!=0:
                    for ele in result['detected_urls']:
                        if result['detected_urls'][0]['positives'] != 0:
                            row = tab2.add_row().cells
                            row[0].text=([ele["url"]])
                            row[1].text=([f'{ele["positives"]}/{ele["total"]}'])
            except KeyError:
                row = tab2.add_row().cells
                row[0].text = (["no info"])
                row[1].text = (["no info"])
            #print('\n\n', tabulate(lt2, headers=["url","score"], tablefmt="orgtbl"))
            print("document created successfully")
            document.save('my_word.docx')
        except Exception as e:
            print(f' unable to create word document| reason:{e}')
    else:
        print(f'status_code is {response.status_code} | reason is {response.reason}')

def valid_ip(ip_name):
    try:
        ip_name=IPAddress(ip_name)
    except AddrFormatError:
        print('invalid ipaddress.provide data as ips')
        return False

    else:
        if ip_name.version in [4,6]:
            if ip_name.is_private():
                print(f'provided {ip_name} is private.don\'t use private ips for reputation')
                return False
            else:return True
        else: return True
#def word(lt1,lt2,lt3):
#try:
 #   document= Document()
  #  document.add_heading("virustotal",level=2)
   # table=document.add_table(rows=2,cols=2)
    #table.style='Table Grid'
 #   row=table.rows[0].cells
  #  row[0].text='ip'
  #  row[1].text='country'
 #   row=table.rows[1].cells
##    row[1].text=
#    table1=document.add_table(rows=1,cols=2)
#    table1.style = 'Table Grid'
#    row = table.rows[0].cells
#    row[0].text = 'last_resolved'
#    row[1].text = 'hostname'
 #   row = table1.add_row().cells
 #   for i in row:
 #       for cell in i.cells:
 #           cell.text=lt3

  ##  document.save('my_word.docx')
#except Exception as e:
   # print(f' unable to create word document| reason:{e}')

def main():
    ip_name=input("enter ip:")
    if valid_ip(ip_name):
        ip_address(ip_name)

main()
