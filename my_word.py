from docx import Document
def word(lt2):
    try:
        document= Document()
        document.add_heading("virustotal",level=2)
        table=document.add_table(rows=20,cols=6)
        table.style='Table Grid'
        row=table.rows[0].cells
        row[0].text='ip'
        row[1].text='country'
        row[2].text='last_resolved'
        row[3].text='hostname'
        row[4].text='url'
        row[5].text='score'
        for i in lt2:
            print(i)
            table.add_row(i)
        print("document created successfully")
        document.save('my_word.docx')
    except Exception as e:
        print(f' unable to create word document| reason:{e}')